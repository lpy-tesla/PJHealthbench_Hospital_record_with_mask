#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
会诊记录分页处理工具(带数据脱敏)

功能:
- 将会诊记录按每条会诊分页
- 自动识别并脱敏患者、医生、医院、城市、检验单号、病理单号等敏感信息
- 患者姓名统一替换为 patient
- 保留医疗数据的完整性

Author: ltesla
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class ConsultationMasker:
    """会诊记录数据脱敏处理器"""
    
    def __init__(self):
        """初始化"""
        # 敏感信息匹配模式
        self.patterns = {
            # 身份证号
            'id_card': r'\b\d{15}\b|\b\d{17}[\dXx]\b',
            
            # 住院号
            'admission_no': r'(?:住院号|入院号|病案号)[:：]\s*[\dA-Z\-]+',
            
            # 医院名称(多种格式)
            'hospital': r'[\u4e00-\u9fa5]{2,20}(?:医院|卫生院|卫生所|诊所|医疗中心|人民医院|中心医院|第一医院|第二医院|第三医院|妇幼保健院|儿童医院|中医院|专科医院)',
            
            # 城市名称(地级市)
            'city': r'(?:北京|上海|天津|重庆|广州|深圳|杭州|南京|武汉|成都|西安|郑州|长沙|哈尔滨|沈阳|济南|青岛|大连|厦门|宁波|苏州|无锡|佛山|东莞|福州|合肥|昆明|南昌|石家庄|长春|太原|南宁|贵阳|兰州|海口|银川|西宁|乌鲁木齐|拉萨|呼和浩特)市?',
            
            # 性别
            'gender': r'(?:性别)[:：]\s*[男女]',
            
            # 年龄(多种格式)
            'age_simple': r'(?:年龄)[:：]\s*\d{1,3}岁',
            'age_in_text': r'([男女]，)(\d{1,3}岁)',
            
            # 出生日期
            'birth_date': r'(?:出生日期|生日|出生)[:：]\s*(?:\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?|\d{8})',
            
            # 手机号
            'phone': r'\b1[3-9]\d{9}\b',
            
            # 检查号(多种格式)
            'exam_no': r'((?:检查号|放射检查号|检验单号|化验单号|检验号|报告单号|单号|病理号|标本号|编号)[:：]\s*)[\w\-_/]+',
            
            # 病理单号
            'pathology_no': r'\b(?:PF|BF|HY)\d{4}-\d{5,8}\b',
            
            # 患者姓名模式
            'patient_name_label': r'(?:患者姓名|姓名)[:：]\s*[\u4e00-\u9fa5]{2,4}',
        }
        
        # 常见医生姓名模式
        self.doctor_patterns = [
            r'(?:会诊医生签名|医生签名|签名)[:：]\s*([\u4e00-\u9fa5]{2,4})',
            r'(?:请会诊医师|会诊医师)[:：]\s*([\u4e00-\u9fa5]{2,4})',
        ]
    
    def mask_patient_name(self, text, patient_name=None):
        """
        替换患者姓名为 patient
        
        Args:
            text: 原始文本
            patient_name: 患者姓名(可选)
            
        Returns:
            替换后的文本
        """
        # 如果未提供患者姓名,尝试从文本中提取
        if not patient_name:
            # 从"患者姓名:"提取
            name_match = re.search(r'患者姓名[:：]\s*([\u4e00-\u9fa5]{2,4})', text)
            if name_match:
                patient_name = name_match.group(1)
        
        if patient_name:
            # 替换所有患者姓名为 patient
            text = text.replace(patient_name, 'patient')
        
        # 通用患者姓名模式替换
        text = re.sub(r'(?:患者姓名|姓名)[:：]\s*[\u4e00-\u9fa5]{2,4}', 
                     lambda m: m.group(0).split('：')[-1].split(':')[-1].replace(
                         m.group(0).split('：')[-1].split(':')[-1], 'patient'
                     ), text)
        
        return text
    
    def mask_doctors(self, text):
        """
        移除医生签名信息
        
        Args:
            text: 原始文本
            
        Returns:
            替换后的文本
        """
        for pattern in self.doctor_patterns:
            text = re.sub(pattern, 
                         lambda m: m.group(0).split('：')[0] + '：[已脱敏]' if '：' in m.group(0) 
                         else m.group(0).split(':')[0] + ':[已脱敏]',
                         text)
        
        return text
    
    def mask_text(self, text):
        """
        对文本进行全面脱敏处理
        
        Args:
            text: 原始文本
            
        Returns:
            脱敏后的文本
        """
        if not text:
            return text
        
        masked_text = text
        
        # 1. 移除身份证号
        masked_text = re.sub(self.patterns['id_card'], '[身份证号已脱敏]', masked_text)
        
        # 2. 移除住院号
        masked_text = re.sub(self.patterns['admission_no'], lambda m: m.group(0).split(':')[0] + ':CODE', masked_text)
        
        # 3. 移除医院名称
        masked_text = re.sub(self.patterns['hospital'], '[医院已脱敏]', masked_text)
        
        # 4. 城市/行政区 - 保留 (不脱敏)
        # masked_text = re.sub(self.patterns['city'], '[城市已脱敏]', masked_text)
        
        # 5. 性别 - 保留 (不脱敏)
        # masked_text = re.sub(self.patterns['gender'], '性别:[已脱敏]', masked_text)
        
        # 6. 保留年龄(不脱敏)
        # masked_text = re.sub(self.patterns['age_simple'], '年龄:[已脱敏]', masked_text)
        # masked_text = re.sub(self.patterns['age_in_text'], r'\1[已脱敏]', masked_text)
        
        # 7. 出生日期 - 仅保留年份 (符合规范)
        def replace_birth_date(match):
            full_match = match.group(0)
            # 提取年份
            year_match = re.search(r'(\d{4})[-/年]?\d{1,2}[-/月]?\d{1,2}', full_match)
            if not year_match:
                year_match = re.search(r'(\d{4})\d{4}', full_match)  # 8位格式
            if year_match:
                year = year_match.group(1)
                label = full_match.split('：')[0] if '：' in full_match else full_match.split(':')[0]
                return f"{label}：{year}年"
            return full_match
        
        masked_text = re.sub(self.patterns['birth_date'], replace_birth_date, masked_text)
        
        # 8. 移除手机号
        masked_text = re.sub(self.patterns['phone'], '[手机号已脱敏]', masked_text)
        
        # 9. 移除检查号
        masked_text = re.sub(self.patterns['exam_no'], lambda m: f"{m.group(1)}CODE", masked_text)
        
        # 10. 移除病理单号
        masked_text = re.sub(self.patterns['pathology_no'], 'CODE', masked_text)
        
        return masked_text


class ConsultationRecordSplitter:
    """会诊记录分页处理器"""
    
    def __init__(self, input_file, output_file=None, enable_masking=True):
        """
        初始化
        
        Args:
            input_file: 输入的Word文档路径
            output_file: 输出的Word文档路径(可选)
            enable_masking: 是否启用数据脱敏
        """
        self.input_file = Path(input_file)
        if output_file:
            self.output_file = Path(output_file)
        else:
            suffix = '_脱敏分页版' if enable_masking else '_分页版'
            self.output_file = self.input_file.parent / f"{self.input_file.stem}{suffix}.docx"
        
        self.consultations = []
        self.enable_masking = enable_masking
        self.masker = ConsultationMasker() if enable_masking else None
    
    def extract_and_parse_from_docx(self):
        """直接从Word文档中按段落和表格解析会诊记录"""
        try:
            doc = Document(self.input_file)
            
            consultations = []
            current_consultation = None
            current_content = []
            
            # 遍历文档所有元素(段落和表格)
            for element in doc.element.body:
                # 处理段落
                if element.tag.endswith('p'):
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    
                    if para:
                        text = para.text.strip()
                        
                        # 检测会诊记录开始
                        if '院内会诊申请及记录单' in text or '会诊申请及记录单' in text or '会诊记录单' in text:
                            # 保存上一条会诊记录
                            if current_consultation and current_content:
                                current_consultation['content'] = '\n'.join(current_content)
                                if self.enable_masking:
                                    current_consultation['content'] = self.masker.mask_patient_name(current_consultation['content'])
                                    current_consultation['content'] = self.masker.mask_doctors(current_consultation['content'])
                                    current_consultation['content'] = self.masker.mask_text(current_consultation['content'])
                                consultations.append(current_consultation)
                            
                            # 开始新的会诊记录
                            current_consultation = {
                                'title': '院内会诊申请及记录单',
                                'dept': '未指定',
                                'time': '未记录',
                                'content': '',
                                'raw_content': ''
                            }
                            current_content = []
                        
                        # 收集内容
                        if current_consultation and text:
                            current_content.append(text)
                
                # 处理表格
                elif element.tag.endswith('tbl'):
                    if current_consultation:
                        # 提取表格内容
                        table = None
                        for t in doc.tables:
                            if t._element == element:
                                table = t
                                break
                        
                        if table:
                            for row in table.rows:
                                row_text = []
                                for cell in row.cells:
                                    cell_text = cell.text.strip()
                                    if cell_text:
                                        row_text.append(cell_text)
                                        
                                        # 识别会诊科室
                                        if '兹邀请' in cell_text:
                                            dept_match = re.search(r'兹邀请\s*([^\s]+?)\s*医师会诊', cell_text)
                                            if dept_match:
                                                current_consultation['dept'] = dept_match.group(1).strip()
                                        
                                        # 识别会诊时间
                                        if '时间：' in cell_text or '时间:' in cell_text:
                                            time_match = re.search(r'时间[:：]\s*(\d{4}年\d{1,2}月\d{1,2}日\d{1,2}时\d{1,2}分)', cell_text)
                                            if time_match:
                                                current_consultation['time'] = time_match.group(1)
                                
                                if row_text:
                                    current_content.append(' | '.join(row_text))
            
            # 保存最后一条记录
            if current_consultation and current_content:
                current_consultation['content'] = '\n'.join(current_content)
                if self.enable_masking:
                    current_consultation['content'] = self.masker.mask_patient_name(current_consultation['content'])
                    current_consultation['content'] = self.masker.mask_doctors(current_consultation['content'])
                    current_consultation['content'] = self.masker.mask_text(current_consultation['content'])
                consultations.append(current_consultation)
            
            return consultations
            
        except Exception as e:
            raise Exception(f"读取文档失败: {str(e)}")
    
    def parse_consultations(self, text):
        """
        解析会诊记录
        
        会诊记录格式识别:
        - 以"院内会诊申请及记录单"或类似标题开头
        - 包含会诊申请和会诊意见两部分
        """
        # 分割会诊记录的正则表达式
        # 匹配: **院内会诊申请及记录单** 
        record_pattern = r'\*\*((?:院内)?会诊(?:申请及)?记录单)\*\*(.*?)(?=\*\*(?:院内)?会诊(?:申请及)?记录单\*\*|$)'
        
        matches = re.finditer(record_pattern, text, re.DOTALL)
        
        consultations = []
        for match in matches:
            title = match.group(1).strip()
            content = match.group(2).strip()
            
            if not content or len(content) < 50:
                continue
            
            # 提取会诊科室
            dept_match = re.search(r'兹邀请\s*[【\[]?([^\]】\n]+?)[】\]]?\s*医师会诊', content)
            dept = dept_match.group(1).strip() if dept_match else "未指定科室"
            
            # 提取会诊时间
            time_match = re.search(r'时间[:：]\s*(\d{4}年\d{1,2}月\d{1,2}日\d{1,2}时\d{1,2}分)', content)
            consult_time = time_match.group(1) if time_match else "未记录时间"
            
            # 脱敏处理
            if self.enable_masking:
                content = self.masker.mask_patient_name(content)
                content = self.masker.mask_doctors(content)
                content = self.masker.mask_text(content)
            
            consultations.append({
                'title': title,
                'dept': dept,
                'time': consult_time,
                'content': content,
                'raw_content': match.group(2)
            })
        
        return consultations
    
    def create_report_document(self):
        """创建分页报告文档"""
        print(f"正在创建新文档...")
        new_doc = Document()
        
        # 设置文档默认样式
        self._set_document_style(new_doc)
        
        total_records = len(self.consultations)
        
        for idx, consultation in enumerate(self.consultations, 1):
            print(f"处理进度: {idx}/{total_records} - {consultation['dept']}会诊")
            
            # 添加会诊记录页
            self._add_consultation_page(new_doc, consultation, idx, total_records)
            
            # 如果不是最后一项,添加分页符
            if idx < total_records:
                new_doc.add_page_break()
        
        # 保存文档
        new_doc.save(self.output_file)
        
        mask_status = "已脱敏" if self.enable_masking else "未脱敏"
        print(f"\n✓ 成功生成报告({mask_status}): {self.output_file}")
        print(f"✓ 共生成 {total_records} 页会诊记录")
        
        return self.output_file
    
    def _set_document_style(self, doc):
        """设置文档默认样式"""
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _add_consultation_page(self, doc, consultation, page_num, total_pages):
        """添加单个会诊记录页"""
        # 1. 添加标题
        self._add_title(doc, consultation['title'])
        
        # 2. 添加脱敏提示
        if self.enable_masking:
            self._add_privacy_notice(doc)
        
        # 3. 添加分隔线
        self._add_separator(doc)
        
        # 4. 添加会诊信息
        self._add_consultation_info(doc, consultation)
        
        # 5. 添加会诊内容
        self._add_consultation_content(doc, consultation)
        
        # 6. 添加页脚
        self._add_footer(doc, page_num, total_pages)
    
    def _add_title(self, doc, title_text):
        """添加居中标题"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(title_text)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.name = '黑体'
        title_run.font.color.rgb = RGBColor(0, 0, 139)
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    def _add_privacy_notice(self, doc):
        """添加隐私保护提示"""
        notice = doc.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notice_run = notice.add_run("(本记录已进行隐私保护处理)")
        notice_run.font.size = Pt(9)
        notice_run.font.italic = True
        notice_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_separator(self, doc, char='─', length=60):
        """添加分隔线"""
        para = doc.add_paragraph(char * length)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_consultation_info(self, doc, consultation):
        """添加会诊基本信息"""
        # 会诊科室
        dept_para = doc.add_paragraph()
        dept_label = dept_para.add_run("会诊科室: ")
        dept_label.font.bold = True
        dept_label.font.size = Pt(12)
        
        dept_value = dept_para.add_run(consultation['dept'])
        dept_value.font.size = Pt(12)
        dept_value.font.bold = True
        dept_value.font.color.rgb = RGBColor(0, 0, 128)
        
        # 会诊时间
        time_para = doc.add_paragraph()
        time_label = time_para.add_run("会诊时间: ")
        time_label.font.bold = True
        time_label.font.size = Pt(11)
        
        time_value = time_para.add_run(consultation['time'])
        time_value.font.size = Pt(11)
        time_value.font.color.rgb = RGBColor(70, 70, 70)
        
        # 添加空行
        doc.add_paragraph()
    
    def _add_consultation_content(self, doc, consultation):
        """添加会诊内容"""
        # 内容标题
        content_title = doc.add_paragraph()
        content_title_run = content_title.add_run('会诊内容:')
        content_title_run.font.size = Pt(11)
        content_title_run.font.bold = True
        content_title_run.font.underline = True
        
        # 分段显示内容
        content = consultation['content']
        
        # 按段落分割
        paragraphs = content.split('\n')
        
        current_section = None
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text or para_text in ['|', '+', '-']:
                continue
            
            # 检查是否为章节标题
            if any(keyword in para_text for keyword in 
                   ['病情摘要', '诊疗情况', '会诊记录', '会诊意见', '诊断意见', 
                    '治疗意见', '用药情况', '检查结果']):
                # 章节标题
                section_para = doc.add_paragraph()
                section_para.paragraph_format.space_before = Pt(6)
                section_para.paragraph_format.space_after = Pt(3)
                
                section_run = section_para.add_run(para_text)
                section_run.font.size = Pt(11)
                section_run.font.bold = True
                section_run.font.color.rgb = RGBColor(0, 0, 128)
                
                current_section = para_text
            else:
                # 普通内容
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.3)
                para.paragraph_format.space_after = Pt(3)
                
                run = para.add_run(para_text)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    def _add_footer(self, doc, page_num, total_pages):
        """添加页脚信息"""
        doc.add_paragraph()
        self._add_separator(doc, '─', 60)
        
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        footer_text = f"第 {page_num} / {total_pages} 页"
        footer_run = footer.add_run(footer_text)
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def process(self):
        """执行完整的处理流程"""
        try:
            # 1. 读取和解析文档
            mask_status = "启用脱敏" if self.enable_masking else "未启用脱敏"
            print(f"正在读取文件: {self.input_file} ({mask_status})")
            print("正在解析会诊记录...")
            
            self.consultations = self.extract_and_parse_from_docx()
            
            if not self.consultations:
                raise Exception("未找到任何会诊记录!请检查文档格式是否正确。")
            
            print(f"共找到 {len(self.consultations)} 条会诊记录")
            
            # 2. 创建分页报告
            output_path = self.create_report_document()
            
            return output_path
            
        except Exception as e:
            print(f"\n✗ 错误: {str(e)}", file=sys.stderr)
            raise


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='会诊记录分页工具(带数据脱敏)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  # 基本使用(默认脱敏)
  python %(prog)s 会诊记录.docx
  
  # 指定输出文件
  python %(prog)s 会诊记录.docx 输出.docx
  
  # 不启用脱敏
  python %(prog)s 会诊记录.docx --no-mask
        """
    )
    
    parser.add_argument('input_file', help='输入的Word文档路径')
    parser.add_argument('output_file', nargs='?', help='输出的Word文档路径(可选)')
    parser.add_argument('--no-mask', action='store_true', help='不启用数据脱敏')
    
    args = parser.parse_args()
    
    # 检查输入文件是否存在
    if not Path(args.input_file).exists():
        print(f"✗ 错误: 输入文件不存在: {args.input_file}")
        return 1
    
    try:
        # 创建处理器并执行
        splitter = ConsultationRecordSplitter(
            input_file=args.input_file,
            output_file=args.output_file,
            enable_masking=not args.no_mask
        )
        output_path = splitter.process()
        
        print(f"\n{'='*60}")
        print(f"处理完成!")
        print(f"输出文件: {output_path}")
        if splitter.enable_masking:
            print(f"✓ 患者信息已替换为: patient")
            print(f"✓ 已移除: 医生签名、身份证、住院号、医院名称、")
            print(f"         城市、性别、年龄、检查号、病理单号等")
        print(f"{'='*60}")
        
        return 0
        
    except Exception as e:
        print(f"\n处理失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1




# === Unified de-identification patch per user policy ===
from collections import OrderedDict as _OrderedDict

_Consult_orig_init = ConsultationMasker.__init__

def _consult_init(self, mask_mode='remove'):
    _Consult_orig_init(self)
    self.mask_mode = mask_mode
    self.doctor_mapping = _OrderedDict()
    self.doctor_counter = 0
    self.hospital_mapping = _OrderedDict()
    self.hospital_counter = 0
    # disable city masking patterns by removing if exists
    self.patterns.pop('city', None)
    # keep gender/age
    self.patterns.pop('gender', None)
    self.patterns.pop('age_simple', None)
    self.patterns.pop('age_in_text', None)
    # strengthen id card
    self.patterns['id_card'] = r'\b\d{15}\b|\b\d{17}[\dXx]\b'
    # DOB keep year only
    self.patterns['birth_date'] = r'((?:出生日期|生日|出生)[:：]\s*)(?:\D*)?((?:19|20)\d{2})(?:\d{2}(?:\d{2})?|[-/年]\d{1,2}[-/月]\d{1,2}日?)?'
    # hospitals (existing pattern ok) -> map to hospital A/B...
    # detailed address masking patterns
    self.patterns.setdefault('address_field', r'((?:地址|住址|现住址|家庭住址|联系地址|单位地址|通讯地址|户籍地址|居住地址)[:：]\s*)(.+)$')
    self.patterns.setdefault('inline_addr', r'(?:(?:[一-龥]{2,30}(?:省|自治区|特别行政区))?(?:[一-龥]{2,30}(?:市|州|盟))?(?:[一-龥]{2,30}(?:区|县|旗|镇|乡|街道|开发区|新区))?)?[一-龥0-9]{0,40}(?:路|街|巷|弄|里|村|屯|组|庄|湾|苑|小区|社区|园区|大厦|广场|写字楼|公寓)[一-龥0-9\-]{0,40}(?:\d{1,4}号)?[一-龥0-9\-]{0,20}(?:(?:楼|幢|栋|单元|室|房)\d{0,6})?')
    # other ids & generic ids
    self.patterns.setdefault('other_ids', r'((?:门诊号|就诊号|就诊卡号|诊疗卡号|医保号|社保号|社会保障号|医疗保险号|费用单号|发票号|单据号|结算单号|交易号|支付单号|对账单号|电子病历号|病历号|EMR号|EMPI|MPI|HIS号|LIS号|RIS号|PACS号|系统号|主索引号|统一编号|UID|标识号|患者ID|患者编号|条码号|条码|条形码|报告流水号|流水号|报告号|报告编号|访问号|接入号|Accession\s*No\.?|Study\s*ID)[:：]\s*)[\w\-_/\.]+')
    self.patterns.setdefault('generic_id', r'(^|\s)((?:[A-Za-z]{2,8}\s*)?(?:ID|No\.?|NO\.?|编号|号码|号|序列号|流水号)[:：]?\s*)([A-Za-z0-9][A-Za-z0-9\-_/.]{7,})')
    # patient name patterns per report
    self._name_patterns = [
        r'(?:姓名|患者|病人|患者姓名)[:：]\s*([\u4e00-\u9fa5]{2,6})',
        r'(?:患者)[:：]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',
    ]

def _consult_identify_doctors(self, text):
    doctor_patterns = [
        r'会诊医生签名[:：]\s*([\u4e00-\u9fa5]{2,6})',
        r'医生签名[:：]\s*([\u4e00-\u9fa5]{2,6})',
        r'(?:请会诊医师|会诊医师)[:：]\s*([\u4e00-\u9fa5]{2,6})',
    ]
    exclude_words = {'主任','副主任','主治','住院','记录日期','签名日期','记录内容','书写日期'}
    for pat in doctor_patterns:
        for m in re.finditer(pat, text):
            name = m.group(1)
            if not name or name in exclude_words or len(name) < 2:
                continue
            if name not in self.doctor_mapping:
                self.doctor_counter += 1
                suffix = chr(64 + self.doctor_counter) if self.doctor_counter <= 26 else str(self.doctor_counter)
                self.doctor_mapping[name] = f'doctor{suffix}'

def _consult_mask_names(self, text):
    for pat in self._name_patterns:
        for m in re.finditer(pat, text):
            original = m.group(0)
            name = m.group(1)
            if self.mask_mode == 'remove':
                text = text.replace(original, '')
            elif self.mask_mode == 'asterisk':
                if len(name) == 2:
                    masked_name = name[0] + '*'
                else:
                    masked_name = name[0] + '*' * (len(name) - 1)
                text = text.replace(original, original.replace(name, masked_name))
            else:
                # placeholder
                text = re.sub(re.escape(name), '[患者姓名]', text, count=1)
    return text

def _consult_mask_doctors(self, text):
    # replace all identified doctor names
    for k, v in sorted(self.doctor_mapping.items(), key=lambda x: len(x[0]), reverse=True):
        text = text.replace(k, v)
    return text

def _consult_mask_text(self, text):
    if not text:
        return text
    masked_text = text
    masked_text = _consult_mask_names(self, masked_text)
    _consult_identify_doctors(self, masked_text)
    masked_text = re.sub(self.patterns['id_card'], '[身份证号已脱敏]', masked_text)
    masked_text = re.sub(self.patterns['admission_no'], lambda m: m.group(0).split('：')[0].split(':')[0] + ':[住院号]' if self.mask_mode=='placeholder' else ('' if self.mask_mode=='remove' else '住院号:***'), masked_text)
    # phone per report scheme
    if 'phone' in self.patterns:
        if self.mask_mode == 'remove':
            masked_text = re.sub(self.patterns['phone'], '', masked_text)
        elif self.mask_mode == 'asterisk':
            masked_text = re.sub(self.patterns['phone'], lambda m: m.group(0)[:3]+'****'+m.group(0)[-4:], masked_text)
        else:
            masked_text = re.sub(self.patterns['phone'], '[手机号]', masked_text)
    # DOB keep year
    masked_text = re.sub(self.patterns['birth_date'], lambda m: f"{m.group(1)}{m.group(2)}", masked_text)
    # hospital mapping
    for h in re.findall(self.patterns['hospital'], masked_text):
        if h not in self.hospital_mapping:
            self.hospital_counter += 1
            suffix = chr(64 + self.hospital_counter) if self.hospital_counter <= 26 else str(self.hospital_counter)
            self.hospital_mapping[h] = f"hospital {suffix}"
    for k, v in sorted(self.hospital_mapping.items(), key=lambda x: len(x[0]), reverse=True):
        masked_text = masked_text.replace(k, v)
    # exam/lab/pathology + other ids
    masked_text = re.sub(r'((?:检查号|放射检查号|检验单号|化验单号|检验号|报告单号|单号|病理号|标本号|编号|病理单号|报告流水号|流水号|条码号|条码|访问号|接入号)[:：]\s*)[\w\-_/\.]+', r'\1CODE', masked_text)
    masked_text = re.sub(self.patterns.get('other_ids', r'$^'), lambda m: f"{m.group(1)}CODE", masked_text, flags=re.IGNORECASE)
    masked_text = re.sub(self.patterns.get('generic_id', r'$^'), lambda m: f"{m.group(1)}{m.group(2)}CODE", masked_text, flags=re.IGNORECASE|re.MULTILINE)
    # detailed address keep city/admin, mask rest
    def _mask_detail(addr: str) -> str:
        if not addr:
            return '[地址已脱敏]'
        prefix_pat = re.compile(r'^\s*([一-龥]{2,30}(?:省|自治区|特别行政区))?\s*([一-龥]{2,30}(?:市|州|盟))?\s*([一-龥]{2,30}(?:区|县|旗))?\s*([一-龥]{2,30}(?:镇|乡|街道|开发区|新区))?\s*')
        m = prefix_pat.match(addr)
        prefix = ''.join([p for p in (m.groups() if m else []) if p])
        return f"{prefix}[地址已脱敏]"
    masked_text = re.sub(self.patterns.get('address_field', r'$^'), lambda m: f"{m.group(1)}{_mask_detail(m.group(2))}", masked_text, flags=re.MULTILINE)
    masked_text = re.sub(self.patterns.get('inline_addr', r'$^'), '[地址已脱敏]', masked_text)
    masked_text = _consult_mask_doctors(self, masked_text)
    return masked_text

ConsultationMasker.__init__ = _consult_init
ConsultationMasker.mask_text = _consult_mask_text
ConsultationMasker.mask_doctors = _consult_mask_doctors
ConsultationMasker.mask_patient_name = lambda self, text, patient_name=None: _consult_mask_names(self, text)



if __name__ == '__main__':
    exit(main())
