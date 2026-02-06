#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
病程记录分页处理工具(统一脱敏方案)

功能:
- 将病程记录按每条记录分页
- 统一的脱敏规则
- 患者姓名统一替换为 patient
- 医生姓名自动识别并替换为 doctorA, doctorB, doctorC...
- 医院名称映射为 hospital A/B/C

Author: ltesla
"""

import re
import sys
from pathlib import Path
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class UnifiedMedicalMasker:
    """统一的医疗数据脱敏处理器"""
    
    def __init__(self):
        """初始化"""
        # 医生姓名映射 (来自medical_record_splitter_modified_v6.py)
        self.doctor_mapping = OrderedDict()
        self.doctor_counter = 0
        
        
        self.patient_names = set()
        # 医院名称映射
        self.hospital_mapping = OrderedDict()
        self.hospital_counter = 0
        
        
        self.title_hospitals = []
        self.main_hospital_label = None# 敏感信息匹配模式
        self.patterns = {
            # 身份证号 - 统一使用medical_report_splitter_with_mask.py的方案(更简洁)
            'id_card': r'\b\d{15}|\d{17}[\dXx]\b',
            
            # 住院号/病案号 - 统一使用medical_report_splitter_with_mask.py的方案
            'admission_no': r'(?:住院号|入院号|病案号)[:：]\s*[\dA-Z\-]+',
            
            # 手机号 - 统一使用medical_report_splitter_with_mask.py的方案
            'phone': r'\b1[3-9]\d{9}\b',
            
            # 医院名称
            'hospital': r'[\u4e00-\u9fa5]{2,20}(?:医院|卫生院|卫生所|诊所|医疗中心|人民医院|中心医院|第一医院|第二医院|第三医院)',
            
            # 性别 - 保留,不脱敏
            'gender': r'(?:性别)[:：]\s*[男女]',
            
            # 年龄 - 保留,不脱敏
            'age_pattern': r'(?:[\u4e00-\u9fa5]+，)([男女]，)(\d{1,3}岁)',
            
            # 出生日期 - 只保留年份
            'birth_date_full': r'(?:出生日期|生日|出生)[:：]\s*(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})日?',
            'birth_date_short': r'(?:出生日期|生日|出生)[:：]\s*(\d{8})',
            
            # 检验/检查/病理单号 - 统一使用modified_v6的方案
            'exam_codes': r'((?:检查号|放射检查号|检验单号|化验单号|检验号|病理单号|病理号|标本号)[:：]\s*)[\w\-_/]+',
            
            # 其他编号 - 统一使用modified_v6的方案
            'other_codes': r'((?:门诊号|就诊卡号|医保卡号|社保卡号|费用单号|电子病历号|条码号|流水号)[:：]\s*)[\w\-_/]+',
            
            # 详细地址 - 改进版,避免误判医学术语
            # 只匹配明确的地址标签开头的内容
            'address_field': r'((?:地址|住址|现住址|家庭住址|联系地址|单位地址|通讯地址|户籍地址|居住地址)[:：]\s*)([^\n]+)$',
        }
        
        # 患者姓名识别模式 - 统一使用medical_report_splitter_with_mask.py的方案
        self.name_patterns = [
            r'(?:姓名|患者|病人)[:：]\s*([\u4e00-\u9fa5]{2,4})',  # 中文姓名
        ]
        
        # 医生姓名识别模式 - 统一使用medical_record_splitter_modified_v6.py的方案
        self.doctor_patterns = [
            r'医生签名[:：]\s*([\u4e00-\u9fa5]{2,4})',
            r'([\u4e00-\u9fa5]{2,4})(?:主任医师|副主任医师|主治医师)(?:查房记录|查房录)',
        ]
        
        # 需要排除的词汇(不是医生姓名)
        self.exclude_doctor_words = [
            '主任', '副主任', '主治', '住院', '请输入', '危重病例', 
            '医师首次', '主任医师', '副主任', '机及随车', '请选择',
            '书写日期', '记录内容', '记录日期', '签名日期'
        ]
        
        # 医学术语白名单 - 这些词不应被当作地址
        self.medical_terms_whitelist = [
            '免疫组化', '辅助分类', '需免疫', '免疫', '组化',
            '病理', '切片', '染色', '标本', '蜡块',
            'S-100', 'SOX-10', 'PNL2', 'MiTF', 'Red-KI-67',
            '克隆号', '正常着色', '阴性', '阳性'
        ]
    
    def identify_patient_names(self, text):
        """从强信号字段中提取患者姓名候选,用于全局精确替换"""
        if not text:
            return
        # 仅依赖字段名/抬头结构信号,不做语义猜测
        pats = [
            r'(?:^|\n|\r)(?:\s*)(?:姓名|患者姓名|病人姓名|患者|病人)[:：\s]+([\u4e00-\u9fa5]{2,4})(?=\s|\n|\r|$|，|,|；|;|\t)',
            r'(?:^|\n|\r)(?:\s*)([\u4e00-\u9fa5]{2,4})(?=\s*(?:男|女)\s*[，,]\s*\d{1,3}岁)',  # 抬头常见: 姓名 男,xx岁
        ]
        for p in pats:
            for m in re.finditer(p, text):
                nm = (m.group(1) or '').strip()
                if 2 <= len(nm) <= 4:
                    self.patient_names.add(nm)

    def identify_doctors(self, text):
        """识别文本中所有的医生姓名"""
        for pattern in self.doctor_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                doctor_name = match.group(1)
                
                # 过滤非姓名词汇
                if doctor_name in self.exclude_doctor_words:
                    continue
                
                # 过滤包含特定词的误识别
                if any(word in doctor_name for word in ['病例', '首次', '记录', '日期']):
                    continue
                
                # 过滤单字姓名
                if len(doctor_name) < 2:
                    continue
                
                if doctor_name not in self.doctor_mapping:
                    self.doctor_counter += 1
                    suffix = chr(64 + self.doctor_counter) if self.doctor_counter <= 26 else str(self.doctor_counter)
                    self.doctor_mapping[doctor_name] = f'doctor{suffix}'
    
    def identify_hospitals(self, text):
        title_text = text
        m = re.search(r'\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}', text)
        if m:
            title_text = text[:m.start()]
        title_matches = re.findall(self.patterns['hospital'], title_text)
        seen = set()
        self.title_hospitals = []
        for h in title_matches:
            if h not in seen:
                seen.add(h)
                self.title_hospitals.append(h)

        all_matches = re.findall(self.patterns['hospital'], text)
        all_seen = set()
        all_hospitals = []
        for h in all_matches:
            if h not in all_seen:
                all_seen.add(h)
                all_hospitals.append(h)

        self.hospital_mapping.clear()
        self.hospital_counter = 0

        for h in self.title_hospitals:
            self.hospital_counter += 1
            suffix = chr(ord('A') + self.hospital_counter - 1) if self.hospital_counter <= 26 else str(self.hospital_counter)
            self.hospital_mapping[h] = f'hospital {suffix}'
        if self.title_hospitals:
            self.main_hospital_label = self.hospital_mapping[self.title_hospitals[0]]

        for h in all_hospitals:
            if h in self.hospital_mapping:
                continue
            self.hospital_counter += 1
            suffix = chr(ord('A') + self.hospital_counter - 1) if self.hospital_counter <= 26 else str(self.hospital_counter)
            self.hospital_mapping[h] = f'hospital {suffix}'
    
    def mask_patient_name(self, text, patient_name=None):
        """替换患者姓名为 patient (统一使用medical_report_splitter方案)"""
        # 如果未提供患者姓名,尝试从文本中提取
        if not patient_name:
            for pattern in self.name_patterns:
                match = re.search(pattern, text)
                if match:
                    patient_name = match.group(1)
                    break
        
        if patient_name:
            text = text.replace(patient_name, 'patient')
        
        # 通用患者模式替换
        text = re.sub(r'(?:姓名|患者)[:：]\s*[\u4e00-\u9fa5]{2,4}', 
                     lambda m: m.group(0).rsplit('：', 1)[0].rsplit(':', 1)[0] + ':patient' 
                     if ':' in m.group(0) else m.group(0).rsplit('：', 1)[0] + '：patient',
                     text)
        
        return text
    
    def mask_doctors(self, text):
        """替换医生姓名为 doctorA, doctorB等"""
        sorted_doctors = sorted(self.doctor_mapping.items(), 
                               key=lambda x: len(x[0]), 
                               reverse=True)
        
        for doctor_name, replacement in sorted_doctors:
            text = text.replace(doctor_name, replacement)
        
        return text
    
    def mask_hospitals(self, text):
        if self.main_hospital_label:
            text = text.replace('我院', self.main_hospital_label)
            text = text.replace('本院', self.main_hospital_label)

        sorted_hospitals = sorted(self.hospital_mapping.items(),
                                  key=lambda x: len(x[0]),
                                  reverse=True)

        for hospital_name, replacement in sorted_hospitals:
            text = text.replace(hospital_name, replacement)

        return text
    
    def mask_birth_date(self, text):
        """处理出生日期 - 只保留年份"""
        # 处理完整日期格式: 1980-01-15 → 1980
        def replace_full_date(match):
            year = match.group(1)
            label = match.group(0).split('：')[0].split(':')[0]
            return f"{label}：{year}年"
        
        text = re.sub(self.patterns['birth_date_full'], replace_full_date, text)
        
        # 处理短格式: 19800115 → 1980
        def replace_short_date(match):
            date_str = match.group(1)
            year = date_str[:4]
            label = match.group(0).split('：')[0].split(':')[0]
            return f"{label}：{year}年"
        
        text = re.sub(self.patterns['birth_date_short'], replace_short_date, text)
        
        return text
    
    def is_medical_context(self, text_before, text_after):
        """判断是否为医学术语上下文,避免误判为地址"""
        # 检查前后文是否包含医学术语
        context = (text_before + text_after).lower()
        
        # 检查是否包含医学术语白名单
        for term in self.medical_terms_whitelist:
            if term.lower() in context:
                return True
        
        # 检查是否包含典型的医学表达
        medical_patterns = [
            r'病理.*?提示', r'免疫.*?化', r'切片', r'染色',
            r'检查.*?示', r'辅助', r'测定', r'检验',
            r'\([+\-]\)', r'阴性|阳性', r'克隆号'
        ]
        
        for pattern in medical_patterns:
            if re.search(pattern, context):
                return True
        
        return False
    
    def mask_address(self, text):
        """处理详细地址 - 改进版,避免误判"""
        # 只处理明确标注为地址的字段
        def replace_address_field(match):
            label = match.group(1)
            addr = match.group(2).strip()
            
            # 提取省市区前缀
            prefix_pat = re.compile(r'^\s*([一-龥]{2,30}(?:省|自治区|特别行政区))?\s*([一-龥]{2,30}(?:市|州|盟))?\s*([一-龥]{2,30}(?:区|县|旗))?\s*')
            m = prefix_pat.match(addr)
            prefix = ''
            if m:
                parts = [p for p in m.groups() if p]
                prefix = ''.join(parts)
            
            return f"{label}{prefix}[详细地址已脱敏]"
        
        # 只替换明确的地址字段
        text = re.sub(self.patterns['address_field'], 
                     replace_address_field, 
                     text, 
                     flags=re.MULTILINE)
        
        return text
    
    def mask_text(self, text):
        """对文本进行统一脱敏处理"""
        if not text:
            return text
        
        masked_text = text
        
        # 额外覆盖：转科/转入转出等表格里常见的“姓名，男/女，xx岁”无字段名写法
        # 仅在紧跟性别与年龄的模式下替换，避免误伤普通名词（如“腮腺区”等）
        masked_text = re.sub(r'(?<![\u4e00-\u9fff])([\u4e00-\u9fff]{2,4})(?=，\s*(男|女)\s*，\s*\d{1,3}岁)', 'patient', masked_text)
        masked_text = re.sub(r'(?<![\u4e00-\u9fff])([\u4e00-\u9fff]{2,4})(?=,\s*(男|女)\s*,\s*\d{1,3}岁)', 'patient', masked_text)
        masked_text = re.sub(r'(?<![\u4e00-\u9fff])([\u4e00-\u9fff]{2,4})(?=\s+(男|女)\s+\d{1,3}岁)', 'patient', masked_text)

        # 1. 移除身份证号
        masked_text = re.sub(self.patterns['id_card'], '[身份证号已脱敏]', masked_text)
        
        # 2. 移除住院号/病案号
        masked_text = re.sub(self.patterns['admission_no'], 
                            lambda m: m.group(0).split('：')[0].split(':')[0] + ':CODE',
                            masked_text)
        
        # 3. 替换医院名称
        masked_text = self.mask_hospitals(masked_text)
        
        # 4. 保留性别 - 不脱敏
        # (不处理)
        
        # 5. 保留年龄 - 不脱敏
        # (不处理)
        
        # 6. 处理出生日期 - 只保留年份
        masked_text = self.mask_birth_date(masked_text)
        
        # 7. 移除手机号
        masked_text = re.sub(self.patterns['phone'], '[手机号已脱敏]', masked_text)
        
        # 8. 处理检验/检查/病理单号
        masked_text = re.sub(self.patterns['exam_codes'], r'\1CODE', masked_text)
        
        # 9. 处理其他编号
        masked_text = re.sub(self.patterns['other_codes'], r'\1CODE', masked_text)
        
        # 10. 处理详细地址 (改进版,避免误判)
        masked_text = self.mask_address(masked_text)
        
        return masked_text


class MedicalRecordSplitter:
    """病程记录分页处理器"""
    
    def __init__(self, input_file, output_file=None, enable_masking=True):
        """
        初始化
        
        Args:
            input_file: 输入的Word文档路径
            output_file: 输出的Word文档路径(可选)
            enable_masking: 是否启用数据脱敏
        """
        self.input_file = Path(input_file)
        if output_file:
            self.output_file = Path(output_file)
        else:
            suffix = '_统一脱敏版' if enable_masking else '_分页版'
            self.output_file = self.input_file.parent / f"{self.input_file.stem}{suffix}.docx"
        
        self.records = []
        self.enable_masking = enable_masking
        self.masker = UnifiedMedicalMasker() if enable_masking else None
    
    def extract_text_from_docx(self):
        """从Word文档中提取文本(段落+表格)"""
        try:
            doc = Document(self.input_file)
            parts = []
            for para in doc.paragraphs:
                if para.text is not None:
                    parts.append(para.text)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        t = (cell.text or '').strip()
                        if t:
                            parts.append(t)
            full_text = chr(10).join(parts)
            return full_text
        except Exception as e:
            raise Exception(f"读取文档失败: {str(e)}")
    
    def parse_medical_records(self, text):
        if self.enable_masking:
            self.masker.identify_doctors(text)
            self.masker.identify_hospitals(text)
            self.masker.identify_patient_names(text)

        dt_pat = re.compile(r'(?m)^\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}')
        kw = [
            '转科（转入）记录','转科（转出）记录','转科(入)记录','转科(出)记录','转科（入）记录','转科（出）记录','危重病例副主任医师查房记录','危重病例主治医师查房记录','危重病例查房记录','主治医师查房记录','副主任医师查房记录','主任医师查房记录','危重病例主任医师查房记录','输血前记录','输血前评估','输血评估','输血前检查','危急值记录','危急值报告','危急值通知','危急值病程记录','术前评估','术前小结','术前讨论','手术风险评估表','术后首次病程记录','术后首次病程','术后记录','术后评估','手术风险评估表','手术记录','麻醉记录','麻醉术前访视','麻醉术前评估','会诊记录','抢救记录','转科记录','转入记录','转出记录','死亡记录','入院记录','出院记录'
        ]
        kw_pat = re.compile(r'(?m)^(?:\s*(?:\d+[\)）\.、]|[一二三四五六七八九十]+[、\.])\s*)?(?:\*\*)?(%s)(?:\*\*)?\s*[:：]?' % '|'.join(map(re.escape, kw)))

        starts = []
        for mm in dt_pat.finditer(text):
            starts.append(mm.start())
        for mm in kw_pat.finditer(text):
            starts.append(mm.start())

        starts = sorted(set(starts))
        if not starts:
            return []

        records = []
        for i, s in enumerate(starts):
            e = starts[i+1] if i+1 < len(starts) else len(text)
            seg = text[s:e].strip()
            if not seg:
                continue

            dt = ''
            content = seg
            mdt = re.match(r'^(\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2})\s*(.*)$', seg, flags=re.DOTALL)
            if mdt:
                dt = mdt.group(1).strip()
                content = mdt.group(2).strip()

            type_match = re.search(r'\*\*([^*]+)\*\*', content)
            if type_match:
                record_type = type_match.group(1).strip()
            else:
                first_line = content.splitlines()[0].strip() if content else ''
                mkw = kw_pat.match(first_line)
                record_type = mkw.group(1).strip() if mkw else '病程记录'

            if self.enable_masking:
                content = self.masker.mask_patient_name(content)
                content = self.masker.mask_doctors(content)
                content = self.masker.mask_text(content)

            records.append({
                'datetime': dt if dt else '无时间戳',
                'type': record_type,
                'content': content,
                'raw_content': seg
            })

                # 合并“仅包含块标题/模板字段”的碎片记录：常见于 转科/危急值/术前术后/输血/风险评估 等
        # 规则：当前块无时间戳且内容极短，且内容基本等于标题（或只包含标题/书写日期等），则并入下一块并继承其时间戳
        merged = []
        i = 0

        def _norm_type(t: str) -> str:
            if not t:
                return ''
            return (t.replace('（', '(').replace('）', ')')
                     .replace('　', ' ').strip())

        def _is_title_only(rec) -> bool:
            t = _norm_type(rec.get('type', ''))
            c = (rec.get('content') or '').strip()
            if not t:
                return False
            c_norm = c.replace('\r', '').strip()
            # 非常短且只包含标题/少量模板词
            if len(c_norm) == 0:
                return True
            # 只包含标题（可能重复一次）
            if _norm_type(c_norm) == t:
                return True
            # 标题 + 书写日期/记录内容 这类模板字段（无实质内容）
            if len(c_norm) <= 120:
                tmp = re.sub(r'\s+', ' ', c_norm)
                tmp_norm = _norm_type(tmp)
                if tmp_norm == t:
                    return True
                if tmp_norm.startswith(t) and all(k in tmp_norm for k in []):
                    return True
                # 仅包含标题 + “书写日期/记录内容/签名”等字段名但无正文
                if tmp_norm.startswith(t) and re.fullmatch(rf"{re.escape(t)}(\s*(书写日期|记录内容|医生签名|医师签名|签名)[:：]?\s*)?", tmp_norm):
                    return True
            return False

        while i < len(records):
            cur = records[i]
            cur_type = _norm_type(cur.get('type', ''))
            if cur.get('datetime') == '无时间戳' and cur_type and _is_title_only(cur) and i + 1 < len(records):
                nxt = records[i + 1]
                # 让下一块继承标题类型；同时避免把标题重复插入内容
                nxt_type = _norm_type(nxt.get('type', ''))
                # 仅当下一块类型是默认“病程记录/无类型”时才覆盖，避免误改已明确类型
                if (not nxt_type) or (nxt_type in ('病程记录', '查房记录', '记录', '病程')):
                    nxt['type'] = cur.get('type', nxt.get('type'))
                # 若下一块内容开头没有标题，则把标题作为首行补进去（但不重复）
                nxt_content = (nxt.get('content') or '').lstrip()
                if cur_type and not _norm_type(nxt_content[:len(cur_type) + 4]).startswith(cur_type):
                    joiner = chr(10)
                    nxt['content'] = (cur_type + joiner + nxt_content).strip()
                i += 1
                continue

            merged.append(cur)
            i += 1

        records = merged


        return records
    def create_report_document(self):
        """创建分页报告文档"""
        print(f"正在创建新文档...")
        new_doc = Document()
        
        # 设置文档默认样式
        self._set_document_style(new_doc)
        
        total_records = len(self.records)
        
        for idx, record in enumerate(self.records, 1):
            print(f"处理进度: {idx}/{total_records} - {record['type']}")
            
            # 添加记录页
            self._add_record_page(new_doc, record, idx, total_records)
            
            # 如果不是最后一项,添加分页符
            if idx < total_records:
                new_doc.add_page_break()
        
        # 保存文档
        new_doc.save(self.output_file)
        
        mask_status = "已脱敏" if self.enable_masking else "未脱敏"
        print(f"\n✓ 成功生成报告({mask_status}): {self.output_file}")
        print(f"✓ 共生成 {total_records} 页病程记录")
        
        if self.enable_masking:
            print(f"\n脱敏映射:")
            if self.masker.doctor_mapping:
                print(f"  医生: {len(self.masker.doctor_mapping)}位")
                for k, v in list(self.masker.doctor_mapping.items())[:5]:
                    print(f"    {k} → {v}")
                if len(self.masker.doctor_mapping) > 5:
                    print(f"    ... 等共{len(self.masker.doctor_mapping)}位医生")
            
            if self.masker.hospital_mapping:
                print(f"  医院: {len(self.masker.hospital_mapping)}家")
                for k, v in self.masker.hospital_mapping.items():
                    print(f"    {k} → {v}")
        
        return self.output_file
    
    def _set_document_style(self, doc):
        """设置文档默认样式"""
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _add_record_page(self, doc, record, page_num, total_pages):
        """添加单个病程记录页"""
        # 1. 添加标题
        self._add_title(doc, "病程记录")
        
        # 2. 添加脱敏提示
        if self.enable_masking:
            self._add_privacy_notice(doc)
        
        # 3. 添加分隔线
        self._add_separator(doc)
        
        # 4. 添加记录信息
        self._add_record_info(doc, record)
        
        # 5. 添加记录内容
        self._add_record_content(doc, record)
        
        # 6. 添加页脚
        self._add_footer(doc, page_num, total_pages)
    
    def _add_title(self, doc, title_text):
        """添加居中标题"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(title_text)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.name = '黑体'
        title_run.font.color.rgb = RGBColor(0, 0, 139)
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    def _add_privacy_notice(self, doc):
        """添加隐私保护提示"""
        notice = doc.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notice_run = notice.add_run("(本记录已进行统一脱敏处理)")
        notice_run.font.size = Pt(9)
        notice_run.font.italic = True
        notice_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_separator(self, doc, char='─', length=60):
        """添加分隔线"""
        para = doc.add_paragraph(char * length)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_record_info(self, doc, record):
        """添加记录基本信息"""
        # 记录时间
        datetime_para = doc.add_paragraph()
        datetime_label = datetime_para.add_run("记录时间: ")
        datetime_label.font.bold = True
        datetime_label.font.size = Pt(11)
        
        datetime_value = datetime_para.add_run(record['datetime'])
        datetime_value.font.size = Pt(11)
        datetime_value.font.color.rgb = RGBColor(70, 70, 70)
        
        # 记录类型
        type_para = doc.add_paragraph()
        type_label = type_para.add_run("记录类型: ")
        type_label.font.bold = True
        type_label.font.size = Pt(11)
        
        type_value = type_para.add_run(record['type'])
        type_value.font.size = Pt(11)
        type_value.font.bold = True
        type_value.font.color.rgb = RGBColor(0, 0, 128)
        
        # 添加空行
        doc.add_paragraph()
    
    def _add_record_content(self, doc, record):
        """添加记录内容"""
        # 内容标题
        content_title = doc.add_paragraph()
        content_title_run = content_title.add_run('记录内容:')
        content_title_run.font.size = Pt(11)
        content_title_run.font.bold = True
        content_title_run.font.underline = True
        
        # 分段显示内容
        content = record.get('content') or ''
        lines = [ln.rstrip() for ln in content.splitlines()]

        subtitle_keywords = {
            '病例特点', '拟诊讨论', '诊断及诊断依据', '鉴别诊断', 'VTE评估', '诊疗计划', '诊疗计划及方案',
            '沟通教育', '辅助检查', '体检、专科检查情况', '体检、专科检查', '当前主要矛盾',
            '解决主要矛盾的途径、措施和方法', '病情分析及处理如下', '术后处理措施及注意事项',
            '补充的病史和体征', '分析讨论', '今日患者病情', '注意事项'
        }

        def is_subtitle_line(t: str) -> bool:
            tt = (t or '').strip()
            if not tt:
                return False
            if tt.endswith('：') or tt.endswith(':'):
                return True
            if tt in subtitle_keywords:
                return True
            for k in subtitle_keywords:
                if tt.startswith(k + '：') or tt.startswith(k + ':'):
                    return True
            if re.match(r'^[一二三四五六七八九十\d]+[、．\.]', tt):
                return True
            return False

        def add_normal_paragraph(text_block: str):
            if not text_block.strip():
                return
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(text_block)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)

        def add_bold_subtitle(line: str):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(line.strip())
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 128)

        buf = []
        for ln in lines + ['']:
            if not ln.strip():
                if buf:
                    add_normal_paragraph("\n".join(buf))
                    buf = []
                continue

            if is_subtitle_line(ln):
                if buf:
                    add_normal_paragraph("\n".join(buf))
                    buf = []
                add_bold_subtitle(ln)
                continue

            buf.append(ln)


    def _add_footer(self, doc, page_num, total_pages):
        """添加页脚信息"""
        doc.add_paragraph()
        self._add_separator(doc, '─', 60)
        
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        footer_text = f"第 {page_num} / {total_pages} 页"
        footer_run = footer.add_run(footer_text)
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def process(self):
        """执行完整的处理流程"""
        try:
            # 1. 读取文档
            mask_status = "启用统一脱敏" if self.enable_masking else "未启用脱敏"
            print(f"正在读取文件: {self.input_file} ({mask_status})")
            text = self.extract_text_from_docx()
            
            # 2. 解析病程记录
            print("正在解析病程记录...")
            self.records = self.parse_medical_records(text)
            
            if not self.records:
                raise Exception("未找到任何病程记录!请检查文档格式是否正确。")
            
            print(f"共找到 {len(self.records)} 条病程记录")
            
            # 3. 创建分页报告
            output_path = self.create_report_document()
            
            return output_path
            
        except Exception as e:
            print(f"\n✗ 错误: {str(e)}", file=sys.stderr)
            raise


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='病程记录分页工具(统一脱敏方案)',
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    parser.add_argument('input_file', help='输入的Word文档路径')
    parser.add_argument('output_file', nargs='?', help='输出的Word文档路径(可选)')
    parser.add_argument('--no-mask', action='store_true', help='不启用数据脱敏')
    
    args = parser.parse_args()
    
    if not Path(args.input_file).exists():
        print(f"✗ 错误: 输入文件不存在: {args.input_file}")
        return 1
    
    try:
        splitter = MedicalRecordSplitter(
            input_file=args.input_file,
            output_file=args.output_file,
            enable_masking=not args.no_mask
        )
        output_path = splitter.process()
        
        print(f"\n{'='*60}")
        print(f"处理完成!")
        print(f"输出文件: {output_path}")
        print(f"{'='*60}")
        
        return 0
        
    except Exception as e:
        print(f"\n处理失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    exit(main())
