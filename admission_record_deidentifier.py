#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
入院录脱敏处理工具
严格遵循《医疗文书脱敏规范（De-identification Specification）》

脱敏规则:
┌─────────────┬──────────────────────┐
│ 信息类型     │ 脱敏方式              │
├─────────────┼──────────────────────┤
│ 患者姓名     │ patient               │
│ 医生姓名     │ doctorA/B/C          │
│ 性别         │ 保留                  │
│ 年龄         │ 保留                  │
│ 出生日期     │ 仅保留年份            │
│ 医院名称     │ hospital A/B/C       │
│ 城市/行政区  │ 保留                  │
│ 详细地址     │ [地址已脱敏]         │
│ 身份证号     │ [身份证号已脱敏]     │
│ 手机号       │ [手机号已脱敏]       │
│ 医疗编号     │ CODE                  │
│ 时间信息     │ 保留                  │
└─────────────┴──────────────────────┘

Author: ltesla
"""

import re
import sys
from pathlib import Path
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


class MedicalDeIdentifier:
    """医疗文书脱敏处理器（符合脱敏规范）"""
    
    def __init__(self):
        """初始化"""
        # 医生映射 (doctorA, doctorB, doctorC...)
        self.doctor_mapping = OrderedDict()
        self.doctor_counter = 0
        
        # 医院映射 (hospital A, hospital B, hospital C...)
        self.hospital_mapping = OrderedDict()
        self.hospital_counter = 0
        
        # 患者姓名
        self.patient_name = None
    
    def extract_patient_name(self, text):
        """从文本中提取患者姓名"""
        # 匹配 "姓名:XXX" 或 "姓 名:XXX"
        patterns = [
            r'姓\s*名[:：]\s*([\u4e00-\u9fa5]{2,4})',
            r'患者[:：]\s*([\u4e00-\u9fa5]{2,4})',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1)
        
        return None
    
    def identify_hospitals(self, text):
        """识别并映射医院名称"""
        # 医院名称后缀
        hospital_suffixes = [
            '临床病理诊断中心', '病理诊断中心', '医疗中心',
            '人民医院', '中心医院', '第一医院', '第二医院', '第三医院',
            '妇幼保健院', '儿童医院', '中医院', '专科医院', 
            '卫生院', '卫生所', '诊所', '医院'
        ]
        
        # 按长度排序(长的优先匹配)
        hospital_suffixes.sort(key=len, reverse=True)
        
        # 构建正则表达式
        suffix_pattern = '|'.join(re.escape(s) for s in hospital_suffixes)
        
        # 匹配医院名称,排除前面的动词
        pattern = rf'(?:于|在|至|到|复习|前往|转入|转出)?\s*([\u4e00-\u9fa5]{{2,20}}(?:{suffix_pattern}))'
        
        matches = re.finditer(pattern, text)
        seen_hospitals = set()
        
        for match in matches:
            hospital_name = match.group(1).strip()
            
            # 清理医院名称前缀
            for prefix in ['于', '在', '至', '到', '复习', '前往', '转入', '转出']:
                if hospital_name.startswith(prefix):
                    hospital_name = hospital_name[len(prefix):].strip()
            
            # 过滤通用词
            if hospital_name in ['我院', '本院', '贵院', '医院']:
                continue
            
            # 过滤过短的名称
            if len(hospital_name) < 5:
                continue
            
            # 避免重复
            if hospital_name in seen_hospitals:
                continue
            
            seen_hospitals.add(hospital_name)
            
            # 添加到映射
            if hospital_name not in self.hospital_mapping:
                self.hospital_counter += 1
                letter = chr(64 + self.hospital_counter)  # A, B, C...
                self.hospital_mapping[hospital_name] = f'hospital {letter}'
    
    def identify_doctors(self, text):
        """识别并映射医生姓名"""
        # 医生签名模式
        patterns = [
            r'医生签名[:：]\s*([\u4e00-\u9fa5]{2,4})',
            r'(?:主治医师|主任医师|副主任医师|住院医师)[:：]\s*([\u4e00-\u9fa5]{2,4})',
            r'(?:医师|医生)[:：]\s*([\u4e00-\u9fa5]{2,4})',
        ]
        
        # 排除词汇
        exclude_words = ['请选择', '请输入', '主任', '副主任', '主治', '住院']
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                doctor_name = match.group(1)
                
                # 过滤
                if doctor_name in exclude_words or len(doctor_name) < 2:
                    continue
                
                # 添加到映射
                if doctor_name not in self.doctor_mapping:
                    self.doctor_counter += 1
                    letter = chr(64 + self.doctor_counter)  # A, B, C...
                    self.doctor_mapping[doctor_name] = f'doctor{letter}'
    
    def mask_patient_name(self, text):
        """患者姓名 → patient"""
        if self.patient_name:
            text = text.replace(self.patient_name, 'patient')
        return text
    
    def mask_hospitals(self, text):
        """医院名称 → hospital A/B/C"""
        # 按名称长度从长到短排序,避免误替换
        sorted_hospitals = sorted(
            self.hospital_mapping.items(),
            key=lambda x: len(x[0]),
            reverse=True
        )
        
        for hospital_name, replacement in sorted_hospitals:
            text = text.replace(hospital_name, replacement)
        
        return text
    
    def mask_doctors(self, text):
        """医生姓名 → doctorA/B/C"""
        sorted_doctors = sorted(
            self.doctor_mapping.items(),
            key=lambda x: len(x[0]),
            reverse=True
        )
        
        for doctor_name, replacement in sorted_doctors:
            text = text.replace(doctor_name, replacement)
        
        return text
    
    def mask_birth_date(self, text):
        """出生日期 → 仅保留年份"""
        # 匹配各种日期格式
        patterns = [
            # 出生日期: YYYY-MM-DD
            (r'出生[日期]*[:：]\s*(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})',
             lambda m: f"出生日期：{m.group(1)}年"),
            
            # 生日: YYYY-MM-DD
            (r'生日[:：]\s*(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})',
             lambda m: f"生日：{m.group(1)}年"),
        ]
        
        for pattern, replacer in patterns:
            text = re.sub(pattern, replacer, text)
        
        return text
    
    def mask_address(self, text):
        """详细地址 → [地址已脱敏] (保留省市区)"""
        # 匹配地址
        address_patterns = [
            r'(?:现住址|家庭住址|住址|地址|户籍地址)[:：]\s*([^\n]+)',
            r'(?:出生地)[:：]\s*([^\n]+)',
        ]
        
        for pattern in address_patterns:
            def replace_address(match):
                full_text = match.group(0)
                address = match.group(1)
                
                # 提取省市区(保留)
                province = re.search(r'([\u4e00-\u9fa5]{2,}省)', address)
                city = re.search(r'([\u4e00-\u9fa5]{2,}市)', address)
                district = re.search(r'([\u4e00-\u9fa5]{2,}(?:区|县))', address)
                
                preserved_parts = []
                if province:
                    preserved_parts.append(province.group(1))
                if city:
                    preserved_parts.append(city.group(1))
                if district:
                    preserved_parts.append(district.group(1))
                
                # 获取标签
                label = full_text.split('：')[0] if '：' in full_text else full_text.split(':')[0]
                
                if preserved_parts:
                    return f"{label}：{' '.join(preserved_parts)} [地址已脱敏]"
                else:
                    return f"{label}：[地址已脱敏]"
            
            text = re.sub(pattern, replace_address, text)
        
        return text
    
    def mask_id_card(self, text):
        """身份证号 → [身份证号已脱敏]"""
        # 18位或15位身份证
        text = re.sub(r'\b\d{15}|\d{17}[\dXx]\b', '[身份证号已脱敏]', text)
        return text
    
    def mask_phone(self, text):
        """手机号 → [手机号已脱敏]"""
        # 11位手机号
        text = re.sub(r'\b1[3-9]\d{9}\b', '[手机号已脱敏]', text)
        return text
    
    def mask_medical_codes(self, text):
        """医疗编号 → CODE"""
        # 住院号、病案号、门诊号等
        medical_code_patterns = [
            (r'(?:住院号|入院号|病案号|门诊号|就诊号)[:：]\s*[\dA-Z\-]+',
             lambda m: m.group(0).split('：')[0] + '：CODE' if '：' in m.group(0) 
                      else m.group(0).split(':')[0] + ':CODE'),
            
            # 检查号、病理号
            (r'(?:检查号|病理号|标本号)[:：]\s*[A-Z0-9_\-]+',
             lambda m: m.group(0).split('：')[0] + '：CODE' if '：' in m.group(0)
                      else m.group(0).split(':')[0] + ':CODE'),
            
            # 病理单号格式: PF2025-08549
            (r'(?:PF|BF|HY)\d{4}-\d{5,8}', 'CODE'),
        ]
        
        for pattern, replacement in medical_code_patterns:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def process_text(self, text):
        """完整的脱敏处理流程"""
        if not text or not text.strip():
            return text
        
        # 按顺序应用所有脱敏规则
        text = self.mask_patient_name(text)
        text = self.mask_hospitals(text)
        text = self.mask_doctors(text)
        text = self.mask_birth_date(text)
        text = self.mask_address(text)
        text = self.mask_id_card(text)
        text = self.mask_phone(text)
        text = self.mask_medical_codes(text)
        
        return text


class AdmissionRecordProcessor:
    """入院录脱敏处理器"""
    
    def __init__(self, input_file, output_file=None):
        """初始化"""
        self.input_file = Path(input_file)
        if output_file:
            self.output_file = Path(output_file)
        else:
            self.output_file = self.input_file.parent / f"{self.input_file.stem}_脱敏版.docx"
        
        self.deidentifier = MedicalDeIdentifier()
    
    def process(self):
        """处理入院录文档"""
        try:
            print(f"\n{'='*70}")
            print(f"入院录脱敏处理工具")
            print(f"遵循《医疗文书脱敏规范（De-identification Specification）》")
            print(f"{'='*70}\n")
            
            print(f"正在读取文件: {self.input_file}")
            doc = Document(self.input_file)
            
            # 第一步: 收集所有文本用于识别
            print("正在分析文档内容...")
            all_text = self._extract_all_text(doc)
            
            # 识别患者姓名
            self.deidentifier.patient_name = self.deidentifier.extract_patient_name(all_text)
            if self.deidentifier.patient_name:
                print(f"  ✓ 识别到患者姓名")
            
            # 识别医院和医生
            self.deidentifier.identify_hospitals(all_text)
            self.deidentifier.identify_doctors(all_text)
            print(f"  ✓ 识别到 {len(self.deidentifier.hospital_mapping)} 个医院")
            print(f"  ✓ 识别到 {len(self.deidentifier.doctor_mapping)} 个医生")
            
            # 第二步: 处理文档内容
            print("\n正在脱敏处理...")
            self._process_paragraphs(doc)
            self._process_tables(doc)
            
            # 第三步: 保存文档
            print(f"\n正在保存文档...")
            doc.save(self.output_file)
            
            # 显示结果
            self._print_results()
            
            return 0
            
        except Exception as e:
            print(f"\n✗ 错误: {str(e)}", file=sys.stderr)
            import traceback
            traceback.print_exc()
            return 1
    
    def _extract_all_text(self, doc):
        """提取文档中的所有文本"""
        all_text = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
        
        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text.append(cell.text)
        
        return '\n'.join(all_text)
    
    def _process_paragraphs(self, doc):
        """处理段落"""
        for para in doc.paragraphs:
            if para.text.strip():
                # 保存原始格式
                original_runs = [(run.text, run.font.size, run.font.bold, 
                                run.font.italic, run.font.name) 
                               for run in para.runs]
                
                # 脱敏处理
                new_text = self.deidentifier.process_text(para.text)
                
                # 清除段落内容
                para.clear()
                
                # 重新添加文本(保持格式)
                run = para.add_run(new_text)
                if original_runs:
                    run.font.size = original_runs[0][1]
                    run.font.bold = original_runs[0][2]
                    run.font.italic = original_runs[0][3]
                    if original_runs[0][4]:
                        run.font.name = original_runs[0][4]
    
    def _process_tables(self, doc):
        """处理表格"""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        # 脱敏处理
                        new_text = self.deidentifier.process_text(cell.text)
                        cell.text = new_text
    
    def _print_results(self):
        """打印处理结果"""
        print(f"\n{'='*70}")
        print(f"✓ 脱敏处理完成!")
        print(f"{'='*70}")
        
        print(f"\n输出文件: {self.output_file}")
        
        # 显示映射关系
        if self.deidentifier.hospital_mapping:
            print(f"\n【医院映射】")
            for original, masked in self.deidentifier.hospital_mapping.items():
                print(f"  {original:30s} → {masked}")
        
        if self.deidentifier.doctor_mapping:
            print(f"\n【医生映射】")
            for original, masked in self.deidentifier.doctor_mapping.items():
                print(f"  {original:30s} → {masked}")
        
        # 显示脱敏规则
        print(f"\n【脱敏规则汇总】")
        rules = [
            ("患者姓名", "patient"),
            ("医生姓名", "doctorA / doctorB / doctorC"),
            ("性别", "保留"),
            ("年龄", "保留"),
            ("出生日期", "仅保留年份"),
            ("医院名称", "hospital A / hospital B / hospital C"),
            ("城市/行政区", "保留"),
            ("详细地址", "[地址已脱敏]"),
            ("身份证号", "[身份证号已脱敏]"),
            ("手机号", "[手机号已脱敏]"),
            ("医疗编号", "CODE"),
            ("时间信息", "保留"),
        ]
        
        for item, rule in rules:
            print(f"  {item:12s} → {rule}")
        
        print(f"\n{'='*70}")


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='入院录脱敏处理工具（遵循医疗文书脱敏规范）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
脱敏规范:
  患者姓名   → patient
  医生姓名   → doctorA / doctorB / doctorC  
  性别       → 保留
  年龄       → 保留
  出生日期   → 仅保留年份
  医院名称   → hospital A / hospital B / hospital C
  城市/行政区 → 保留
  详细地址   → [地址已脱敏]
  身份证号   → [身份证号已脱敏]
  手机号     → [手机号已脱敏]
  医疗编号   → CODE
  时间信息   → 保留

使用示例:
  python %(prog)s 入院录.docx
  python %(prog)s 入院录.docx -o 输出.docx
        """
    )
    
    parser.add_argument('input_file', help='输入的入院录Word文档')
    parser.add_argument('-o', '--output', dest='output_file', 
                       help='输出文件路径(可选)')
    
    args = parser.parse_args()
    
    # 检查输入文件
    if not Path(args.input_file).exists():
        print(f"✗ 错误: 输入文件不存在: {args.input_file}")
        return 1
    
    # 创建处理器并执行
    processor = AdmissionRecordProcessor(args.input_file, args.output_file)
    return processor.process()


if __name__ == '__main__':
    exit(main())
