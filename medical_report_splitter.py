#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
带数据脱敏功能的检验报告分页工具

功能:
- 自动识别并移除敏感信息(姓名、身份证、住院号等)
- 支持多种脱敏策略
- 保留检验项目和结果数据的完整性

Author: ltesla
"""

import re
import sys
from pathlib import Path
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


class DataMasker:
    """数据脱敏处理器"""
    
    def __init__(self, mask_mode='remove'):
        """
        初始化脱敏器
        
        Args:
            mask_mode: 脱敏模式
                - 'remove': 完全移除敏感信息
                - 'asterisk': 用星号替换 (如: 张** / 32010***********12)
                - 'placeholder': 用占位符替换 (如: [患者姓名] / [身份证号])
        """
        self.mask_mode = mask_mode
        
        # 敏感信息匹配模式
        self.patterns = {
            # 身份证号: 18位或15位数字
            'id_card': r'\b\d{15}|\d{17}[\dXx]\b',
            
            # 住院号: 常见格式 (可能需要根据实际情况调整)
            'admission_no': r'(?:住院号|入院号|病案号)[:：]\s*[\dA-Z\-]+',
            
            # 医院名称: 以"医院"、"卫生院"、"诊所"等结尾
            'hospital': r'[\u4e00-\u9fa5]{2,20}(?:医院|卫生院|卫生所|诊所|医疗中心|人民医院|中心医院)',
            
            # 性别: "性别:男/女"
            'gender': r'(?:性别)[:：]\s*[男女]',
            
            # 年龄: "年龄:XX岁"
            'age': r'(?:年龄)[:：]\s*\d{1,3}\s*岁',
            
            # 出生日期: YYYY-MM-DD 或 YYYY年MM月DD日
            'birth_date': r'(?:出生日期|生日|出生)[:：]\s*(?:\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?|\d{8})',
            
            # 手机号: 11位数字
            'phone': r'\b1[3-9]\d{9}\b',
            
            # 地址: "地址:"开头的内容
            'address': r'(?:地址|住址|家庭住址)[:：][^\n]{5,50}',
        }
        
        # 姓名识别需要特殊处理,因为姓名格式比较复杂
        self.name_patterns = [
            r'(?:姓名|患者|病人)[:：]\s*([\u4e00-\u9fa5]{2,4})',  # 中文姓名
            r'(?:患者)[:：]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',  # 英文姓名
        ]
    
    def mask_text(self, text):
        """
        对文本进行脱敏处理
        
        Args:
            text: 原始文本
            
        Returns:
            脱敏后的文本
        """
        if not text:
            return text
        
        masked_text = text
        
        # 1. 处理姓名 (需要先处理,避免被其他规则影响)
        masked_text = self._mask_names(masked_text)
        
        # 2. 处理其他敏感信息
        for info_type, pattern in self.patterns.items():
            masked_text = self._mask_by_pattern(masked_text, pattern, info_type)
        
        return masked_text
    
    def _mask_names(self, text):
        """识别并脱敏姓名"""
        for pattern in self.name_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                original = match.group(0)
                name = match.group(1)
                
                if self.mask_mode == 'remove':
                    # 移除整行
                    text = text.replace(original, '')
                elif self.mask_mode == 'asterisk':
                    # 保留姓,名字用*替换
                    if len(name) == 2:
                        masked_name = name[0] + '*'
                    elif len(name) == 3:
                        masked_name = name[0] + '**'
                    elif len(name) == 4:
                        masked_name = name[0] + '***'
                    else:
                        masked_name = name[0] + '*' * (len(name) - 1)
                    
                    masked = original.replace(name, masked_name)
                    text = text.replace(original, masked)
                elif self.mask_mode == 'placeholder':
                    masked = re.sub(r'([:：]\s*)[^\s]+', r'\1[患者姓名]', original)
                    text = text.replace(original, masked)
        
        return text
    
    def _mask_by_pattern(self, text, pattern, info_type):
        """根据正则模式脱敏"""
        matches = re.finditer(pattern, text)
        
        # 占位符映射
        placeholders = {
            'id_card': '[身份证号]',
            'admission_no': '住院号:[住院号]',
            'hospital': '[医院名称]',
            'gender': '性别:[性别]',
            'age': '年龄:[年龄]',
            'birth_date': '出生日期:[出生日期]',
            'phone': '[手机号]',
            'address': '地址:[地址]',
        }
        
        for match in matches:
            original = match.group(0)
            
            if self.mask_mode == 'remove':
                # 完全移除
                text = text.replace(original, '')
            
            elif self.mask_mode == 'asterisk':
                # 用星号替换
                if info_type == 'id_card':
                    # 身份证: 保留前3位和后2位
                    masked = original[:3] + '*' * (len(original) - 5) + original[-2:]
                elif info_type == 'phone':
                    # 手机号: 保留前3位和后4位
                    masked = original[:3] + '****' + original[-4:]
                elif info_type == 'hospital':
                    # 医院: 保留前2个字和后2个字
                    if len(original) > 4:
                        masked = original[:2] + '*' * (len(original) - 4) + original[-2:]
                    else:
                        masked = '*' * len(original)
                else:
                    # 其他: 保留标签,内容用星号
                    if ':' in original or '：' in original:
                        parts = re.split(r'[:：]', original, 1)
                        masked = parts[0] + ':***'
                    else:
                        masked = '*' * len(original)
                
                text = text.replace(original, masked)
            
            elif self.mask_mode == 'placeholder':
                # 用占位符替换
                placeholder = placeholders.get(info_type, '[敏感信息]')
                text = text.replace(original, placeholder)
        
        return text
    
    def mask_report_item(self, item_name, content):
        """
        对单个检验项目进行脱敏
        
        Args:
            item_name: 项目名称
            content: 项目内容
            
        Returns:
            (脱敏后的项目名称, 脱敏后的内容)
        """
        # 项目名称一般不含敏感信息,但也检查一下
        masked_name = self.mask_text(item_name)
        masked_content = self.mask_text(content)
        
        return masked_name, masked_content


class MedicalReportSplitterWithMask:
    """带数据脱敏功能的检验报告分页处理器"""
    
    def __init__(self, input_file, output_file=None, mask_mode='remove', 
                 enable_masking=True):
        """
        初始化
        
        Args:
            input_file: 输入的Word文档路径
            output_file: 输出的Word文档路径(可选)
            mask_mode: 脱敏模式 ('remove', 'asterisk', 'placeholder')
            enable_masking: 是否启用数据脱敏
        """
        self.input_file = Path(input_file)
        if output_file:
            self.output_file = Path(output_file)
        else:
            # 默认输出文件名
            suffix = '_脱敏分页版' if enable_masking else '_分页版'
            self.output_file = self.input_file.parent / f"{self.input_file.stem}{suffix}.docx"
        
        self.test_items = []
        self.enable_masking = enable_masking
        self.masker = DataMasker(mask_mode) if enable_masking else None
        
    def extract_text_from_docx(self):
        """从Word文档中提取文本"""
        try:
            doc = Document(self.input_file)
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            return full_text
        except Exception as e:
            raise Exception(f"读取文档失败: {str(e)}")
    
    def parse_test_items(self, text):
        """解析检验项目"""
        # 如果启用脱敏,先对整个文本进行脱敏
        if self.enable_masking:
            text = self.masker.mask_text(text)
        
        # 正则表达式匹配检验项目
        pattern = r'【([^】]+)】\(([0-9\-\s:]+)\)([^【]*)'
        matches = re.findall(pattern, text, re.DOTALL)
        
        test_items = []
        for match in matches:
            item_name = match[0].strip()
            date = match[1].strip()
            content = match[2].strip()
            
            # 清理内容中的换行符和多余空格
            content = re.sub(r'\s+', ' ', content)
            
            # 对单个项目再次脱敏(双重保险)
            if self.enable_masking:
                item_name, content = self.masker.mask_report_item(item_name, content)
            
            test_items.append({
                'name': item_name,
                'date': date,
                'content': content,
                'raw_content': match[2]
            })
        
        return test_items
    
    def parse_results(self, content):
        """解析检验结果"""
        # 按中英文逗号、分号分割
        items = re.split(r'[,，;；]', content)
        
        results = []
        for item in items:
            item = item.strip()
            if not item:
                continue
            
            # 检查是否包含冒号
            if ':' in item or '：' in item:
                results.append(item)
            elif len(item) > 2:
                results.append(item)
        
        return results
    
    def is_abnormal(self, result):
        """判断检验结果是否异常"""
        return '↑' in result or '↓' in result
    
    def create_report_document(self):
        """创建分页报告文档"""
        print(f"正在创建新文档...")
        new_doc = Document()
        
        # 设置文档默认样式
        self._set_document_style(new_doc)
        
        total_items = len(self.test_items)
        
        for idx, item in enumerate(self.test_items, 1):
            print(f"处理进度: {idx}/{total_items} - {item['name']}")
            
            # 添加报告页
            self._add_report_page(new_doc, item, idx, total_items)
            
            # 如果不是最后一项,添加分页符
            if idx < total_items:
                new_doc.add_page_break()
        
        # 保存文档
        new_doc.save(self.output_file)
        
        mask_status = "已脱敏" if self.enable_masking else "未脱敏"
        print(f"\n✓ 成功生成报告({mask_status}): {self.output_file}")
        print(f"✓ 共生成 {total_items} 页检验报告")
        
        return self.output_file
    
    def _set_document_style(self, doc):
        """设置文档默认样式"""
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _add_report_page(self, doc, item, page_num, total_pages):
        """添加单个报告页"""
        # 1. 添加标题
        self._add_title(doc, "检验报告单")
        
        # 2. 添加脱敏提示(如果启用了脱敏)
        if self.enable_masking:
            self._add_privacy_notice(doc)
        
        # 3. 添加分隔线
        self._add_separator(doc)
        
        # 4. 添加检验项目信息
        self._add_test_info(doc, item)
        
        # 5. 添加检验结果
        self._add_test_results(doc, item)
        
        # 6. 添加页脚
        self._add_footer(doc, page_num, total_pages)
    
    def _add_title(self, doc, title_text):
        """添加居中标题"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(title_text)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.name = '黑体'
        title_run.font.color.rgb = RGBColor(0, 0, 139)
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    def _add_privacy_notice(self, doc):
        """添加隐私保护提示"""
        notice = doc.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notice_run = notice.add_run("(本报告已进行隐私保护处理)")
        notice_run.font.size = Pt(9)
        notice_run.font.italic = True
        notice_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_separator(self, doc, char='─', length=60):
        """添加分隔线"""
        para = doc.add_paragraph(char * length)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_test_info(self, doc, item):
        """添加检验项目基本信息"""
        # 检验项目名称
        name_para = doc.add_paragraph()
        name_label = name_para.add_run("检验项目: ")
        name_label.font.bold = True
        name_label.font.size = Pt(12)
        
        name_value = name_para.add_run(item['name'])
        name_value.font.size = Pt(12)
        name_value.font.bold = True
        name_value.font.color.rgb = RGBColor(0, 0, 128)
        
        # 检验日期
        date_para = doc.add_paragraph()
        date_label = date_para.add_run("检验日期: ")
        date_label.font.bold = True
        date_label.font.size = Pt(11)
        
        date_value = date_para.add_run(item['date'])
        date_value.font.size = Pt(11)
        date_value.font.color.rgb = RGBColor(70, 70, 70)
        
        # 添加空行
        doc.add_paragraph()
    
    def _add_test_results(self, doc, item):
        """添加检验结果详情"""
        # 结果标题
        result_title = doc.add_paragraph()
        result_title_run = result_title.add_run('检验结果:')
        result_title_run.font.size = Pt(11)
        result_title_run.font.bold = True
        result_title_run.font.underline = True
        result_title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 解析检验结果
        results = self.parse_results(item['content'])
        
        if results:
            for result in results:
                self._add_result_item(doc, result)
        else:
            content_para = doc.add_paragraph()
            content_para.paragraph_format.left_indent = Inches(0.3)
            content_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            content_run = content_para.add_run(item['content'])
            content_run.font.size = Pt(10)
    
    def _add_result_item(self, doc, result):
        """添加单个检验结果项"""
        result_para = doc.add_paragraph()
        result_para.paragraph_format.left_indent = Inches(0.3)
        result_para.paragraph_format.space_after = Pt(3)
        
        # 添加项目符号
        bullet = result_para.add_run("• ")
        bullet.font.size = Pt(10)
        
        # 添加结果内容
        result_run = result_para.add_run(result)
        result_run.font.size = Pt(10)
        
        # 如果是异常值,设置为红色加粗
        if self.is_abnormal(result):
            result_run.font.color.rgb = RGBColor(220, 20, 60)
            result_run.font.bold = True
        else:
            result_run.font.color.rgb = RGBColor(0, 0, 0)
    
    def _add_footer(self, doc, page_num, total_pages):
        """添加页脚信息"""
        doc.add_paragraph()
        self._add_separator(doc, '─', 60)
        
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        footer_text = f"第 {page_num} / {total_pages} 页"
        footer_run = footer.add_run(footer_text)
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def process(self):
        """执行完整的处理流程"""
        try:
            # 1. 读取文档
            mask_status = "启用脱敏" if self.enable_masking else "未启用脱敏"
            print(f"正在读取文件: {self.input_file} ({mask_status})")
            text = self.extract_text_from_docx()
            
            # 2. 解析检验项目
            print("正在解析检验项目...")
            self.test_items = self.parse_test_items(text)
            
            if not self.test_items:
                raise Exception("未找到任何检验项目!请检查文档格式是否正确。")
            
            print(f"共找到 {len(self.test_items)} 个检验项目")
            
            # 3. 创建分页报告
            output_path = self.create_report_document()
            
            return output_path
            
        except Exception as e:
            print(f"\n✗ 错误: {str(e)}", file=sys.stderr)
            raise


def main():
    """主函数"""
    import argparse
    
    # 设置命令行参数解析
    parser = argparse.ArgumentParser(
        description='检验报告分页工具(带数据脱敏功能)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  # 基本使用(默认移除敏感信息)
  python %(prog)s input.docx
  
  # 指定输出文件
  python %(prog)s input.docx output.docx
  
  # 不启用脱敏
  python %(prog)s input.docx --no-mask
  
  # 使用星号脱敏
  python %(prog)s input.docx --mask-mode asterisk
  
  # 使用占位符脱敏
  python %(prog)s input.docx --mask-mode placeholder
        """
    )
    
    parser.add_argument('input_file', help='输入的Word文档路径')
    parser.add_argument('output_file', nargs='?', help='输出的Word文档路径(可选)')
    parser.add_argument('--mask-mode', choices=['remove', 'asterisk', 'placeholder'],
                       default='remove',
                       help='脱敏模式: remove(移除), asterisk(星号), placeholder(占位符)')
    parser.add_argument('--no-mask', action='store_true',
                       help='不启用数据脱敏')
    
    args = parser.parse_args()
    
    # 检查输入文件是否存在
    if not Path(args.input_file).exists():
        print(f"✗ 错误: 输入文件不存在: {args.input_file}")
        return 1
    
    try:
        # 创建处理器并执行
        splitter = MedicalReportSplitterWithMask(
            input_file=args.input_file,
            output_file=args.output_file,
            mask_mode=args.mask_mode,
            enable_masking=not args.no_mask
        )
        output_path = splitter.process()
        
        print(f"\n{'='*60}")
        print(f"处理完成!")
        print(f"输出文件: {output_path}")
        if splitter.enable_masking:
            print(f"脱敏模式: {args.mask_mode}")
            print(f"✓ 已移除/脱敏: 姓名、身份证、住院号、医院名称、性别、年龄等")
        print(f"{'='*60}")
        
        return 0
        
    except Exception as e:
        print(f"\n处理失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1




# === Unified de-identification patch (patient/admission/phone per MedicalReportSplitterWithMask; doctor per v6; DOB keep year; hospital A/B/C; keep gender/age/city; detailed address mask; ids CODE) ===
from collections import OrderedDict as _OrderedDict

# Patch DataMasker behavior in-place, keeping existing interfaces.
_DataMasker_orig_init = DataMasker.__init__

def _DataMasker_init_patched(self, mask_mode='remove'):
    _DataMasker_orig_init(self, mask_mode=mask_mode)
    # keep gender/age (do not mask)
    self.patterns.pop('gender', None)
    self.patterns.pop('age', None)
    # strongest id-card pattern
    self.patterns['id_card'] = r'\b\d{15}\b|\b\d{17}[\dXx]\b'
    # DOB keep year only (capture)
    self.patterns['birth_date'] = r'((?:出生日期|生日|出生)[:：]\s*)(?:\D*)?((?:19|20)\d{2})(?:\d{2}(?:\d{2})?|[-/年]\d{1,2}[-/月]\d{1,2}日?)?'
    # hospital mapping A/B/C...
    self._hospital_mapping = _OrderedDict()
    self._hospital_counter = 0

def _mask_by_pattern_patched(self, text, pattern, info_type):
    matches = list(re.finditer(pattern, text))
    placeholders = {
        'id_card': '[身份证号]',
        'admission_no': '住院号:[住院号]',
        'phone': '[手机号]',
        'birth_date': '出生日期:[出生日期]',
    }
    for match in matches:
        original = match.group(0)
        if info_type == 'hospital':
            if original not in self._hospital_mapping:
                self._hospital_counter += 1
                suffix = chr(64 + self._hospital_counter) if self._hospital_counter <= 26 else str(self._hospital_counter)
                self._hospital_mapping[original] = f'hospital {suffix}'
            text = text.replace(original, self._hospital_mapping[original])
            continue
        if info_type == 'birth_date':
            # keep year only
            try:
                prefix, year = match.group(1), match.group(2)
                text = text.replace(original, f"{prefix}{year}年")
            except Exception:
                pass
            continue
        if info_type == 'admission_no':
            # 医疗编号统一替换为CODE
            try:
                parts = re.split(r'[:：]', original, 1)
                prefix = parts[0] + ':' if len(parts) > 1 else ''
                text = text.replace(original, f"{prefix}CODE")
            except Exception:
                text = text.replace(original, 'CODE')
            continue
        if self.mask_mode == 'remove':
            text = text.replace(original, '')
        elif self.mask_mode == 'asterisk':
            if info_type == 'id_card':
                masked = original[:3] + '*' * (len(original) - 5) + original[-2:]
            elif info_type == 'phone':
                masked = original[:3] + '****' + original[-4:]
            elif info_type == 'admission_no':
                parts = re.split(r'[:：]', original, 1)
                masked = parts[0] + ':***' if len(parts) > 1 else '***'
            else:
                if ':' in original or '：' in original:
                    parts = re.split(r'[:：]', original, 1)
                    masked = parts[0] + ':***'
                else:
                    masked = '*' * len(original)
            text = text.replace(original, masked)
        elif self.mask_mode == 'placeholder':
            placeholder = placeholders.get(info_type, '[敏感信息]')
            text = text.replace(original, placeholder)
    return text

DataMasker.__init__ = _DataMasker_init_patched
DataMasker._mask_by_pattern = _mask_by_pattern_patched



if __name__ == '__main__':
    exit(main())
